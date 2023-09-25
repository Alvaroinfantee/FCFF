import streamlit as st
from docx import Document
from io import BytesIO
import base64

# Function to generate MS Word report with Assumptions
def generate_report(project_name, comments, assumptions, cost_of_equity, cost_of_debt, wacc, fcff_list, npv, price_per_share):
    doc = Document()
    doc.add_heading(f'Financial Report: {project_name}', 0)
    doc.add_heading('Assumptions', level=1)
    doc.add_paragraph(assumptions)
    if comments:
        doc.add_heading('Initial Comments', level=1)
        doc.add_paragraph(comments)
    doc.add_heading('Cost of Equity', level=1)
    doc.add_paragraph(f'The Cost of Equity is: {cost_of_equity * 100:.2f}%')
    doc.add_heading('Cost of Debt', level=1)
    doc.add_paragraph(f'The After-Tax Cost of Debt is: {cost_of_debt * 100:.2f}%')
    doc.add_heading('WACC', level=1)
    doc.add_paragraph(f'The WACC is: {wacc * 100:.2f}%')
    doc.add_heading('FCFF', level=1)
    for i, fcff in enumerate(fcff_list, 1):
        doc.add_paragraph(f'Year {i}: ${fcff:,.2f}')
    doc.add_heading('NPV', level=1)
    doc.add_paragraph(f'The NPV is: ${npv:,.2f}')
    doc.add_heading('Price Per Share', level=1)
    doc.add_paragraph(f'The Price Per Share is: ${price_per_share:,.2f}')
    if comments:
        doc.add_heading('Closing Comments', level=1)
        doc.add_paragraph(comments)
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def get_download_link(file, filename):
    """Generate a link to download the file."""
    filedata = file.read()
    b64 = base64.b64encode(filedata).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download Report</a>'

# Streamlit App
st.title("NPV, WACC, FCFF, and Price Per Share Calculator")

project_name = st.text_input("Project or Company Name:", "Project A")
comments = st.text_area("Comments (Optional):", "")
assumptions = st.text_area("Assumptions (Optional):", "")

# Financial Calculations
risk_free_rate = st.number_input("Risk-Free Rate (%)", min_value=0.0, value=1.0) / 100
beta = st.number_input("Beta", min_value=0.0, value=1.0)
equity_risk_premium = st.number_input("Equity Risk Premium (%)", min_value=0.0, value=6.0) / 100
cost_of_equity = risk_free_rate + beta * equity_risk_premium

cost_of_debt_raw = st.number_input("Yield to Maturity or Interest Rate (%)", min_value=0.0, value=5.0) / 100
tax_rate = st.number_input("Tax Rate (%)", min_value=0.0, value=30.0) / 100
cost_of_debt = cost_of_debt_raw * (1 - tax_rate)

equity_portion = st.number_input("Market Value of Equity ($)", min_value=0.0, value=500000.0)
debt_portion = st.number_input("Market Value of Debt ($)", min_value=0.0, value=500000.0)
total_value = equity_portion + debt_portion
wacc = (equity_portion / total_value * cost_of_equity) + (debt_portion / total_value * cost_of_debt)

years = st.number_input("Number of Years:", min_value=1, value=5)
fcff_list = []
for i in range(years):
    st.subheader(f"Year {i + 1}")
    operating_income = st.number_input(f"Operating Income Year {i + 1} ($)", min_value=0.0, value=100000.0)
    taxes = st.number_input(f"Taxes Year {i + 1} ($)", min_value=0.0, value=30000.0)
    depreciation = st.number_input(f"Depreciation & Amortization Year {i + 1} ($)", min_value=0.0, value=20000.0)
    capex = st.number_input(f"Capital Expenditures & Acquisitions Year {i + 1} ($)", min_value=0.0, value=15000.0)
    change_in_wc = st.number_input(f"Change in non-cash Working Capital Year {i + 1} ($)", min_value=0.0, value=5000.0)
    fcff = operating_income - taxes + depreciation - capex - change_in_wc
    fcff_list.append(fcff)

initial_investment = st.number_input("Initial Investment ($):", min_value=0.0, value=100000.0)
npv = -initial_investment
for t in range(years):
    npv += fcff_list[t] / ((1 + wacc) ** (t + 1))

num_shares = st.number_input("Number of Outstanding Shares:", min_value=1, value=100000)
equity_value = npv - debt_portion
price_per_share = equity_value / num_shares if num_shares != 0 else 0

if st.button("Generate Report"):
    file_obj = generate_report(project_name, comments, assumptions, cost_of_equity, cost_of_debt, wacc, fcff_list, npv, price_per_share)
    download_link = get_download_link(file_obj, f"Financial_Report_{project_name.replace(' ', '_')}.docx")
    st.markdown(download_link, unsafe_allow_html=True)
